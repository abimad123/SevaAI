import os
import time
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from pydantic import BaseModel

router = APIRouter()


@router.post("/analyze")
async def analyze_document(file: UploadFile = File(...), document_id: str = Form(...)):
    """Analyze an uploaded document using AI"""
    try:
        content = await file.read()
        text = extract_text(content, file.filename, file.content_type)

        if not text.strip():
            return {
                "summary": "Could not extract text from this document.",
                "key_information": [],
                "missing_information": ["Document text could not be extracted"],
                "compliance_issues": [],
                "suggestions": ["Please ensure document is not scanned/image-only"],
                "confidence": 0.1,
            }

        # Ingest into vector store
        try:
            from app.services.rag import add_to_vectorstore
            from langchain.schema import Document as LCDocument
            doc = LCDocument(
                page_content=text,
                metadata={"source": file.filename, "document_id": document_id, "type": "uploaded_document"},
            )
            add_to_vectorstore([doc])
        except Exception:
            pass

        # Analyze with LLM
        try:
            from app.services.rag import get_llm
            from langchain.prompts import ChatPromptTemplate

            llm = get_llm(temperature=0.1)
            prompt = ChatPromptTemplate.from_messages([
                ("system", """You are a document analysis expert for Indian NGO sector. 
Analyze the provided document and return a JSON response with these exact keys:
- summary: Brief summary of the document (2-3 sentences)
- key_information: List of key facts/data points as [{key: string, value: string}]
- missing_information: List of missing information that should be present
- compliance_issues: List of any compliance or regulatory issues found
- suggestions: List of actionable suggestions for improvement
- confidence: Float between 0-1 indicating analysis confidence

Be specific to Indian NGO regulations, compliance requirements, and government scheme documentation."""),
                ("human", f"Document filename: {file.filename}\n\nDocument content:\n{text[:3000]}"),
            ])

            import json
            result = llm.invoke(prompt.format_messages())
            content_str = result.content
            # Try to parse JSON from response
            try:
                if "```json" in content_str:
                    content_str = content_str.split("```json")[1].split("```")[0]
                elif "```" in content_str:
                    content_str = content_str.split("```")[1].split("```")[0]
                analysis = json.loads(content_str.strip())
            except json.JSONDecodeError:
                analysis = {
                    "summary": content_str[:500],
                    "key_information": [],
                    "missing_information": [],
                    "compliance_issues": [],
                    "suggestions": [],
                    "confidence": 0.6,
                }
            return analysis
        except Exception as llm_error:
            return {
                "summary": f"Document '{file.filename}' uploaded successfully. AI analysis requires OpenAI API key to be configured.",
                "key_information": [{"key": "File Name", "value": file.filename}, {"key": "File Size", "value": f"{len(content)} bytes"}],
                "missing_information": [],
                "compliance_issues": [],
                "suggestions": ["Configure OPENAI_API_KEY in ai-service/.env to enable full AI analysis"],
                "confidence": 0.0,
            }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def extract_text(content: bytes, filename: str, content_type: str) -> str:
    """Extract text from various document formats"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf" or content_type == "application/pdf":
            import pypdf
            import io
            reader = pypdf.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in ["doc", "docx"]:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs)
        elif ext == "txt" or "text" in (content_type or ""):
            return content.decode("utf-8", errors="ignore")
        else:
            return content.decode("utf-8", errors="ignore")
    except Exception:
        return content.decode("utf-8", errors="ignore")
